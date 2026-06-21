from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from savt import __app_name__, __version__
from savt.chapter_reviews import CHECK_LABELS
from savt.models import AuditReport
from savt.report_builder import findings_dataframe_rows

LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "oia_uccuyo_logo.jpg"
UCCUYO_GREEN = RGBColor(0x06, 0x49, 0x2F)
UCCUYO_GREEN_DARK = RGBColor(0x03, 0x48, 0x2E)
UCCUYO_ORANGE = RGBColor(0xE1, 0x7D, 0x16)
UCCUYO_BG_SOFT = "EEF4F0"
INSTITUTION_LINE = "Universidad Católica de Cuyo · Observatorio de Inteligencia Artificial"


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, side: str, color: str, size: str = "12") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _styled_run(paragraph, text: str, *, bold: bool = False, size: int = 11, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_cover_header(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    logo_cell, text_cell = table.rows[0].cells[0], table.rows[0].cells[1]
    _set_cell_shading(logo_cell, UCCUYO_BG_SOFT)
    _set_cell_shading(text_cell, "06492F")
    _set_cell_border(text_cell, "left", "E17D16", "24")

    if LOGO_PATH.exists():
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))

    title_p = text_cell.paragraphs[0]
    _styled_run(title_p, "SAVT", bold=True, size=24, color=RGBColor(255, 255, 255))
    subtitle_p = text_cell.add_paragraph()
    _styled_run(
        subtitle_p,
        f"{__app_name__} · v{__version__}",
        size=11,
        color=RGBColor(0xE8, 0xF3, 0xED),
    )
    desc_p = text_cell.add_paragraph()
    _styled_run(
        desc_p,
        "Sistema de auditoría de tesis y trabajos finales. Analiza estructura, coherencia "
        "y consistencia del documento, con observaciones y recomendaciones para su mejora académica.",
        size=10,
        color=RGBColor(0xF4, 0xFA, 0xF7),
    )
    inst_p = text_cell.add_paragraph()
    _styled_run(
        inst_p,
        INSTITUTION_LINE.upper(),
        size=8,
        color=RGBColor(0xCF, 0xE3, 0xD8),
    )

    doc.add_paragraph("")


def _add_running_header(doc: Document, filename: str) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.is_linked_to_previous = False
    section.header.is_linked_to_previous = False

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    _remove_table_borders(header_table)
    logo_cell, text_cell = header_table.rows[0].cells[0], header_table.rows[0].cells[1]
    _set_cell_shading(text_cell, UCCUYO_BG_SOFT)

    if LOGO_PATH.exists():
        logo_p = logo_cell.paragraphs[0]
        logo_p.add_run().add_picture(str(LOGO_PATH), width=Inches(0.45))

    text_p = text_cell.paragraphs[0]
    _styled_run(text_p, "SAVT", bold=True, size=10, color=UCCUYO_GREEN)
    _styled_run(text_p, "  ·  Informe de auditoría académica  ·  ", size=9, color=UCCUYO_GREEN_DARK)
    _styled_run(text_p, filename, size=9, color=RGBColor(0x33, 0x33, 0x33))

    bottom = header.add_paragraph()
    bottom.paragraph_format.space_after = Pt(6)
    p_pr = bottom._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom_edge = OxmlElement("w:bottom")
    bottom_edge.set(qn("w:val"), "single")
    bottom_edge.set(qn("w:sz"), "8")
    bottom_edge.set(qn("w:color"), "06492F")
    p_bdr.append(bottom_edge)
    p_pr.append(p_bdr)


def _style_document_headings(doc: Document) -> None:
    for level, size, color in ((1, 14, UCCUYO_GREEN), (2, 12, UCCUYO_GREEN_DARK)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(f"{bold_prefix} ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def _add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)


def build_report_docx(report: AuditReport, dashboard: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    _style_document_headings(doc)
    _add_running_header(doc, report.filename)
    _add_cover_header(doc)

    title = doc.add_heading("Informe de auditoría académica", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Versión del sistema: {__version__}\n").italic = True
    meta.add_run(f"Documento analizado: {report.filename}\n").italic = True
    meta.add_run(f"Fecha del informe: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").italic = True

    doc.add_paragraph("")

    _add_heading(doc, "1. Síntesis (ICAI)", 1)
    _add_label_paragraph(doc, "ICAI", f"{dashboard['icai']}/100 — {dashboard['icai_interpretation']}")
    _add_label_paragraph(
        doc,
        "Estado general",
        f"{dashboard['readiness']}",
    )
    _add_label_paragraph(doc, "Motivo principal", dashboard["main_reason"])
    _add_label_paragraph(doc, "Errores críticos", str(dashboard["errors"]))
    _add_label_paragraph(doc, "Advertencias", str(dashboard["warnings"]))
    _add_label_paragraph(doc, "Páginas estimadas", str(report.page_estimate))
    _add_label_paragraph(doc, "Referencias detectadas", str(len(report.bibliography)))

    _add_heading(doc, "2. Checklist previo a la entrega", 1)
    checklist = dashboard.get("checklist", {})
    _add_label_paragraph(doc, "Estado del checklist", checklist.get("status", "—"))
    for item in checklist.get("items", []):
        if item.get("ok"):
            mark = "Conforme"
        elif item.get("partial"):
            mark = "Parcialmente conforme"
        else:
            mark = "No conforme"
        _add_bullet(doc, f"{item['label']} — {mark}")
    pending_reviews = [r for r in dashboard.get("chapter_reviews") or [] if not r.get("ok")]
    if pending_reviews:
        doc.add_paragraph(
            "El detalle de qué falta, por qué importa y cómo corregir figura en "
            "«Apartados con observaciones» (sección 9)."
        )

    _add_heading(doc, "3. Advertencias detectadas", 1)
    warnings = dashboard.get("warnings_list") or []
    if not warnings:
        doc.add_paragraph("No se registraron advertencias ni errores críticos prioritarios.")
    for idx, item in enumerate(warnings, start=1):
        _add_heading(doc, f"3.{idx} {item['title']}", 2)
        _add_label_paragraph(doc, "Gravedad", item.get("gravity", ""))
        doc.add_paragraph(item.get("detail", ""))
        if item.get("why"):
            _add_label_paragraph(doc, "Por qué importa", item["why"])
        if item.get("how_to_fix"):
            _add_label_paragraph(doc, "Cómo corregir", item["how_to_fix"])
        for row in item.get("detail_items") or []:
            line = f"{row.get('label', '')}: {row.get('value', '')}"
            if row.get("extra"):
                line += f" ({row['extra']})"
            _add_bullet(doc, line)

    _add_heading(doc, "4. Evaluación por SAVT", 1)
    jury = dashboard.get("jury", {})
    doc.add_paragraph("Fortalezas principales:")
    for idx, item in enumerate(jury.get("strengths") or ["—"], start=1):
        _add_bullet(doc, f"{idx}. {item}")
    doc.add_paragraph("Debilidades principales:")
    for idx, item in enumerate(jury.get("weaknesses") or ["—"], start=1):
        _add_bullet(doc, f"{idx}. {item}")
    _add_label_paragraph(
        doc,
        "Probabilidad estimada de aprobación",
        jury.get("approval_probability", "—"),
    )

    _add_label_paragraph(
        doc,
        "Probabilidad estimada de aprobación",
        jury.get("approval_probability", "—"),
    )
    if jury.get("disclaimer"):
        doc.add_paragraph(jury["disclaimer"]).italic = True

    _add_heading(doc, "5. Normativa institucional", 1)
    _add_label_paragraph(doc, "Perfil", dashboard.get("profile_label", "—"))
    formal = dashboard.get("formal_dashboard") or {}
    for key, label in [
        ("portada_completa", "Portada completa"),
        ("indice", "Índice general"),
        ("indice_figuras", "Índice de figuras"),
        ("indice_tablas", "Índice de tablas"),
        ("palabras_clave", "Palabras clave"),
    ]:
        if key in formal:
            _add_bullet(doc, f"{label}: {'Sí' if formal[key] else 'No detectado'}")
    if formal.get("abstract_words"):
        _add_label_paragraph(doc, "Resumen (palabras)", str(formal["abstract_words"]))

    _add_heading(doc, "6. Integridad académica", 1)
    integrity = dashboard.get("integrity_dashboard") or {}
    if integrity.get("disclaimer"):
        doc.add_paragraph(integrity["disclaimer"])
    idx = integrity.get("similarity_index")
    if idx is not None:
        _add_label_paragraph(doc, "Índice de similitud externo", f"{idx:.1f}%")
    else:
        doc.add_paragraph("Sin reporte de similitud externo proporcionado.")

    _add_heading(doc, "7. Ética de investigación", 1)
    ethics = dashboard.get("ethics_dashboard") or {}
    for item in ethics.get("checklist") or []:
        mark = "Sí" if item.get("found") else "No detectado"
        _add_bullet(doc, f"{item.get('label', '')}: {mark}")

    _add_heading(doc, "8. Profundidad académica", 1)
    content = dashboard.get("content_dashboard") or {}
    orig = dashboard.get("originality_dashboard") or {}
    _add_label_paragraph(doc, "Palabras totales (cuerpo)", str(content.get("total_body_words", 0)))
    _add_label_paragraph(doc, "Palabras bibliografía", str(content.get("bibliography_words", 0)))
    for item in content.get("sections") or []:
        _add_bullet(
            doc,
            f"{item.get('title', '—')}: {item.get('words', 0)} palabras "
            f"({item.get('percent_label', '—')} del cuerpo)",
        )
    section_depth = content.get("section_depth") or []
    if section_depth:
        doc.add_paragraph("Profundidad por apartado canónico:")
        for item in section_depth:
            result_part = ""
            if item.get("result_markers"):
                result_part = f", ind. hallazgos {item['result_markers']}"
            _add_bullet(
                doc,
                f"{item.get('title', '—')}: {item.get('words', 0)} palabras, "
                f"{item.get('citation_density', 0)} citas/100 pal., "
                f"{item.get('critical_markers', 0)} marcadores críticos{result_part} "
                f"— {item.get('depth_label', '—')}",
            )
    _add_label_paragraph(doc, "Índice proxy originalidad", f"{orig.get('score_proxy', 0)}/100")

    _add_heading(doc, "9. Apartados con observaciones", 1)
    pending_reviews = [r for r in dashboard.get("chapter_reviews") or [] if not r.get("ok")]
    if not pending_reviews:
        doc.add_paragraph(
            "Todos los apartados evaluados cumplen los criterios detectados automáticamente."
        )
    for review in pending_reviews:
        status = "Parcialmente conforme" if review.get("partial") else "No conforme"
        _add_heading(doc, f"{review['title']} — {status}", 2)
        doc.add_paragraph(review.get("summary", ""))
        missing = (review.get("missing") or []) + (review.get("partial_items") or [])
        if missing:
            doc.add_paragraph("Qué falta o no está claro:")
            for label in missing:
                _add_bullet(doc, CHECK_LABELS.get(label, label))
        if review.get("why"):
            _add_label_paragraph(doc, "Por qué importa", review["why"])
        if review.get("how_to_fix"):
            _add_label_paragraph(doc, "Cómo corregir", review["how_to_fix"])

    _add_heading(doc, "10. Bibliografía", 1)
    bib = dashboard.get("bibliography_dashboard", {})
    lines = [
        f"Estilo: {bib.get('style', '—')}",
        f"Total referencias: {bib.get('total_refs', 0)}",
        f"Citas en texto: {bib.get('citations_found', 0)}",
        f"Citas no emparejadas: {bib.get('unmatched_citations', 0)}",
        f"Fuera del período metodológico: {bib.get('out_of_period', 0)}",
        f"Posiblemente ajenas al tema: {bib.get('possibly_off_topic', 0)}",
        f"Cobertura: {bib.get('coverage', '—')}",
    ]
    for line in lines:
        _add_bullet(doc, line)

    details = bib.get("details") or {}
    unmatched = details.get("unmatched_apa") or []
    if unmatched:
        doc.add_paragraph("Citas APA no emparejadas:")
        for entry in unmatched[:20]:
            cites = ", ".join(entry.get("citations_in_text") or [])
            _add_bullet(doc, f"{cites} (clave: {entry.get('key', '')})")

    _add_heading(doc, "11. Detalle de hallazgos", 1)
    rows = findings_dataframe_rows(report)
    if rows:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Área", "Severidad", "Hallazgo", "Qué significa"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for row in rows[:40]:
            cells = table.add_row().cells
            cells[0].text = str(row.get("Área", ""))[:40]
            cells[1].text = str(row.get("Severidad", ""))[:25]
            cells[2].text = str(row.get("Hallazgo", ""))[:120]
            cells[3].text = str(row.get("Qué significa", ""))[:200]
        if len(rows) > 40:
            doc.add_paragraph(f"(Se omitieron {len(rows) - 40} hallazgos adicionales en la tabla.)")
    else:
        doc.add_paragraph("No se registraron hallazgos adicionales.")

    doc.add_paragraph("")
    footer = doc.add_paragraph(
        "Este informe fue generado automáticamente por SAVT. "
        "No sustituye la evaluación del director de tesis ni del jurado examinador."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
